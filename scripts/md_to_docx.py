from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"C:\0_code\10_read")
SRC = BASE / "proposal.md"
OUT = BASE / "proposal.docx"


def set_east_asian_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_style(style, font_name, size_pt, bold=False, color=None, line_spacing=1.5, space_after=0):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)


def add_run_with_bold(paragraph, text, base_font="SimSun", latin_font="Times New Roman", size=12):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        run.font.name = latin_font
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def make_docx():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    configure_style(styles["Normal"], "SimSun", 12, line_spacing=1.5, space_after=0)
    configure_style(styles["Heading 1"], "SimHei", 16, bold=True, color=(0, 0, 0), line_spacing=1.5, space_after=6)
    configure_style(styles["Heading 2"], "SimHei", 14, bold=True, color=(0, 0, 0), line_spacing=1.5, space_after=3)
    configure_style(styles["Heading 3"], "SimHei", 12, bold=True, color=(0, 0, 0), line_spacing=1.5, space_after=3)

    first_heading = True
    in_refs = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            first_heading = False
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            in_refs = title == "参考文献"
            p = doc.add_paragraph(title, style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.space_after = Pt(3)
            add_run_with_bold(p, line, size=10.5)
            continue

        p = doc.add_paragraph(style="Normal")
        if re.match(r"^(姓名|单位|学号|专业|邮箱)", line):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith("关键词"):
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p.paragraph_format.first_line_indent = Cm(0.74)
        add_run_with_bold(p, line, size=12)

    add_page_number(section.footer.paragraphs[0])

    core_props = doc.core_properties
    core_props.title = "AI 技术革新的辩证法：从“替代劳动”到“重组劳动”"
    core_props.subject = "自然辩证法读书笔记"
    core_props.author = ""

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(make_docx())
